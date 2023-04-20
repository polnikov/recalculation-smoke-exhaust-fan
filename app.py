import os
import sys
import re
import json
import platform
import docx
from datetime import datetime

from docx.shared import Cm, Pt, Mm
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PySide6.QtCore import QSize, Qt, QTimer
from PySide6.QtGui import QColor, QFont, QIcon, QAction
from PySide6.QtWidgets import (
    QApplication,
    QMainWindow,
    QMessageBox,
    QPushButton,
    QLabel,
    QGridLayout,
    QTableWidget,
    QVBoxLayout,
    QGroupBox,
    QWidget,
    QTableWidgetItem,
    QScrollArea,
    QSpacerItem,
    QFileDialog,
)

from constants import CONSTANTS


basedir = os.path.dirname(__file__)


try:
    from ctypes import windll  # Only exists on Windows.
    myappid = 'akudjatechnology.Recalculation-of-the-smoke-exhaust-fan.1.0.0'
    windll.shell32.SetCurrentProcessExplicitAppUserModelID(myappid)
except ImportError:
    pass

now = datetime.now()
version = now.strftime("%Y.%m.%d")


class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.app_title = f'{CONSTANTS.APP_TITLE}_{version}'
        self.setWindowTitle(self.app_title)
        self.groupbox_count = 0
        self.current_file_path = None
        self.box_style = 'QGroupBox::title { color: blue; }'
        
        self.auto_save_timer = QTimer()
        self.auto_save_timer.timeout.connect(self.auto_save)
        self.auto_save_timer.start(300_000) # 5 minutes in milliseconds

        menubar = self.menuBar()
        file_menu = menubar.addMenu(CONSTANTS.MENU[0])

        manual_action = QAction(CONSTANTS.MENU[1], self)
        menubar.addAction(manual_action)

        about_action = QAction(CONSTANTS.MENU[2], self)
        menubar.addAction(about_action)

        open_action = QAction(CONSTANTS.FILE_SUBMENU[0], self)
        open_action.setIcon(QIcon(os.path.join(basedir, 'open.png')))
        open_action.setShortcut("Ctrl+O")
        file_menu.addAction(open_action)
        file_menu.addSeparator()
        save_action = QAction(CONSTANTS.FILE_SUBMENU[1], self)
        save_action.setIcon(QIcon(os.path.join(basedir, 'save.png')))
        save_action.setShortcut("Ctrl+S")
        file_menu.addAction(save_action)

        save_as_action = QAction(CONSTANTS.FILE_SUBMENU[2], self)
        save_as_action.setIcon(QIcon(os.path.join(basedir, 'save_as.png')))
        save_as_action.setShortcut("Ctrl+Shift+S")
        file_menu.addAction(save_as_action)
        file_menu.addSeparator()

        export_action = QAction(CONSTANTS.FILE_SUBMENU[3], self)
        export_action.setIcon(QIcon(os.path.join(basedir, 'export.png')))
        file_menu.addAction(export_action)

        open_action.triggered.connect(self.open)
        save_action.triggered.connect(self.save)
        save_as_action.triggered.connect(self.save_as)
        export_action.triggered.connect(self.export)
        about_action.triggered.connect(self.show_about)
        manual_action.triggered.connect(self.open_manual)

        menubar.setStyleSheet('font-family: Consolas; font-size: 11px;')

        self.widget = QWidget(self)
        self.setCentralWidget(self.widget)
        self.layout = QVBoxLayout(self.widget)
        self.layout.addWidget(self.create_content())

        self.showMaximized()


    def create_content(self) -> object:
        _widget = QWidget()
        _widget.setStyleSheet('background-color: #FFFFFF; border: 0')
        _layout = QVBoxLayout(_widget)

        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)

        _layout.addWidget(self.create_board(), alignment=Qt.AlignmentFlag.AlignTop)
        _layout.addWidget(scroll_area)

        scroll_widget = QWidget(scroll_area)
        scroll_area.setWidget(scroll_widget)

        self.box_tab1 = QVBoxLayout(scroll_widget)
        box_tab1 = self.box_tab1
        box_tab1.addStretch(1)
        box_tab1.addWidget(self.create_table_1(), alignment=Qt.AlignmentFlag.AlignCenter)
        box_tab1.addWidget(self.create_table_2(), alignment=Qt.AlignmentFlag.AlignCenter)
        box_tab1.addWidget(self.create_default_table(), alignment=Qt.AlignmentFlag.AlignCenter)
        box_tab1.addStretch(1)
        return _widget


    def create_board(self) -> object:
        _widget = QWidget()
        _widget.setStyleSheet('background-color: white; ')
        _layout = QGridLayout()
        self.board = _layout

        labels = CONSTANTS.BOARD.LABELS
        for i in range(len(labels)):
            label = QLabel(labels[i])
            label.setStyleSheet('QLabel { border-radius: 20px; background-color: #EFEFEF; }')
            label.setFixedHeight(50)
            font = QFont('Consolas', 12)
            label.setFont(font)
            _layout.addWidget(label, 0, i)
            match i:
                case 0:
                    label.setFixedWidth(890)
                case 2:
                    label.setFixedWidth(150)
                    label.setAlignment(Qt.AlignmentFlag.AlignCenter)
                case 1 | 4 | 5:
                    label.setFixedWidth(60)
                    label.setAlignment(Qt.AlignmentFlag.AlignCenter)
                    if i == 1:
                        label.setText('-')
                case 3:
                    label.setFixedWidth(100)
                    label.setAlignment(Qt.AlignmentFlag.AlignCenter)
                    label.setStyleSheet('QLabel { border-radius: 20px; background-color: #99CCFF; }')
            if i == 5:
                label.setToolTip(CONSTANTS.BOARD.NUM_TABLES_TOOLTIP)
                label.setText('1')
                label.setStyleSheet('QLabel { border-radius: 20px; background-color: #EFEFEF; color: red }')

        self.copy_button = QPushButton()
        copy_button = self.copy_button
        copy_button.setFixedSize(50, 50)
        copy_button.setIcon(QIcon(os.path.join(basedir, 'copy.png')))
        copy_button.setStyleSheet('QPushButton { border-radius: 20px; background-color: #EFEFEF; } QPushButton:hover { border: 2px solid grey; background-color: #99CCFF }')
        copy_button.setToolTip(CONSTANTS.BUTTONS.TOOLTIPS[0])

        copy_button.clicked.connect(self.copy_table)

        self.add_button = QPushButton()
        add_button = self.add_button
        add_button.setFixedSize(50, 50)
        add_button.setIcon(QIcon(os.path.join(basedir, 'add.png')))
        add_button.setStyleSheet('QPushButton { border-radius: 20px; background-color: #EFEFEF; } QPushButton:hover { border: 2px solid grey; background-color: #99FF99 }')
        add_button.setToolTip(CONSTANTS.BUTTONS.TOOLTIPS[1])

        add_button.clicked.connect(self.add_table)

        self.delete_button = QPushButton()
        delete_button = self.delete_button
        delete_button.setFixedSize(50, 50)
        delete_button.setIcon(QIcon(os.path.join(basedir, 'delete.png')))
        delete_button.setStyleSheet('QPushButton { border-radius: 20px; background-color: #EFEFEF; } QPushButton:hover { border: 2px solid grey; background-color: #FF9999 }')
        delete_button.setToolTip(CONSTANTS.BUTTONS.TOOLTIPS[2])

        delete_button.clicked.connect(self.delete_table)

        _layout.addWidget(copy_button, 0, 6)
        _layout.addWidget(delete_button, 0, 7)
        _layout.addWidget(add_button, 0, 8)

        _widget.setLayout(_layout)
        return _widget


    def create_table_1(self) -> object:
        t_rows = CONSTANTS.TABLE1.ROWS
        t_cols = CONSTANTS.TABLE1.COLUMNS

        self.table_1 = QTableWidget(t_rows, t_cols)
        table = self.table_1
        table.setObjectName(CONSTANTS.TABLE1.NAME)
        table.horizontalHeader().setVisible(False)
        table.verticalHeader().setVisible(False)

        for col in range(5):
            table.setColumnWidth(col, CONSTANTS.TABLE_COLUMN_WIDTH.get(col))

        # столбец для ввода
        for row in range(t_rows):
            item = QTableWidgetItem()
            font = QFont('Consolas', 12)
            item.setFont(font)
            table.setItem(row, 3, item)
            if row in CONSTANTS.TABLE1.EDITABLE_ROWS:  # editable
                table.item(row, 3).setBackground(QColor(204, 255, 204))
            else:
                table.item(row, 3).setBackground(QColor(239, 239, 239))
                table.item(row, 3).setFlags(Qt.ItemFlag.ItemIsEnabled)
            table.item(row, 3).setTextAlignment(Qt.AlignmentFlag.AlignCenter)

        for n, row in enumerate(CONSTANTS.TABLE1.EDITABLE_ROWS):
            table.item(row, 3).setToolTip(f'<span style="font-family: Consolas; color: red">{CONSTANTS.TABLE1.INPUTS_TOOLTIPS[n]}</span>')

        # заголовок
        for row in range(t_rows):
            if row in CONSTANTS.TABLE1.SPAN_ROWS:
                table.setSpan(row, 0, 2, 1)

        for row in CONSTANTS.TABLE1.HEADER_ROWS:
            self._set_value_in_cell(row, 0, CONSTANTS.TABLE1.HEADER, table)

        for row in range(t_rows):
            self._set_value_in_cell(row, 1, CONSTANTS.TABLE1.FORMULAS, table)
            self._set_value_in_cell(row, 2, CONSTANTS.TABLE1.SYMBOLS, table)
            self._set_value_in_cell(row, 4, CONSTANTS.TABLE1.UNITS, table)
            table.setRowHeight(row, CONSTANTS.TABLE_ROW_HEIGHT)

        table.itemChanged.connect(self.validate_input_data_in_tables)
        table.itemChanged.connect(self.calculate_Gdpn)
        table.itemChanged.connect(self.calculate_Psn)
        table.itemChanged.connect(self.calculate_Lambda_n)
        table.itemChanged.connect(self.calculate_result)
        table.cellChanged.connect(self.calculate_Tsm0)
        table.cellChanged.connect(self.calculate_Ta)
        table.cellChanged.connect(self.calculate_Tv)
        table.cellChanged.connect(self.calculate_density_a)
        table.cellChanged.connect(self.calculate_density_v)
        table.cellChanged.connect(self.calculate_density_sm)
        table.cellChanged.connect(self.calculate_Psa)
        table.cellChanged.connect(self.calculate_pressure_after_Psa)
        table.cellChanged.connect(self.calculate_Ga)
        table.cellChanged.connect(self._set_Ta_in_board)

        table.setStyleSheet('QTableWidget { border: 0px solid grey; font-family: Consolas; }')
        table.setMinimumWidth(1562)
        table.setMinimumHeight(402)
        return table


    def create_table_2(self) -> object:
        t_rows = CONSTANTS.TABLE2.ROWS
        t_cols = CONSTANTS.TABLE2.COLUMNS

        self.table_2 = QTableWidget(t_rows, t_cols)
        table = self.table_2
        table.setObjectName(CONSTANTS.TABLE2.NAME)
        table.horizontalHeader().setVisible(False)
        table.verticalHeader().setVisible(False)

        for col in range(5):
            table.setColumnWidth(col, CONSTANTS.TABLE_COLUMN_WIDTH.get(col))

        # столбец для ввода
        for row in range(t_rows):
            item = QTableWidgetItem()
            font = QFont('Consolas', 12)
            item.setFont(font)
            table.setItem(row, 3, item)
            if row == 3:  # editable
                table.item(row, 3).setBackground(QColor(204, 255, 204))
            else:
                table.item(row, 3).setBackground(QColor(239, 239, 239))
                table.item(row, 3).setFlags(Qt.ItemFlag.ItemIsEnabled)
            table.item(row, 3).setTextAlignment(Qt.AlignmentFlag.AlignCenter)

        table.item(3, 3).setToolTip(f'<span style="font-family: Consolas; color: red">{CONSTANTS.TABLE2.INPUT_TOOLTIP}</span>')

        # заголовок
        for row in range(t_rows):
            self._set_value_in_cell(row, 0, CONSTANTS.TABLE2.HEADER, table)
            self._set_value_in_cell(row, 1, CONSTANTS.TABLE2.FORMULAS, table)
            self._set_value_in_cell(row, 2, CONSTANTS.TABLE2.SYMBOLS, table)
            self._set_value_in_cell(row, 4, CONSTANTS.TABLE2.UNITS, table)
            table.setRowHeight(row, CONSTANTS.TABLE_ROW_HEIGHT)

        table.itemChanged.connect(self.validate_input_data_in_tables)
        table.itemChanged.connect(self.calculate_Gn)
        table.itemChanged.connect(self.calculate_Psn)
        table.itemChanged.connect(self.calculate_Lambda_n)
        table.itemChanged.connect(self.calculate_result)
        table.cellChanged.connect(self.calculate_Ga)
        table.cellChanged.connect(self.calculate_Psa)
        table.cellChanged.connect(self.calculate_pressure_after_Psa)

        table.setStyleSheet('QTableWidget { border: 0px solid grey; font-family: Consolas; }')
        table.setMinimumWidth(1562)
        table.setMinimumHeight(202)
        return table


    def create_default_table(self) -> object:
        self.groupbox_count += 1
        if self.groupbox_count == 1:
            _box = QGroupBox('Участок до вентилятора')
        else:
            _box = QGroupBox(f'Участок {self.groupbox_count-1}')

        _box.setStyleSheet(self.box_style)
        _box.setFont(QFont('Consolas', 12))
        _layout = QVBoxLayout()
        spacer = QSpacerItem(1500, 10)
        _layout.addSpacerItem(spacer)

        t_rows = CONSTANTS.DEFAULT_TABLE.ROWS
        t_cols = CONSTANTS.DEFAULT_TABLE.COLUMNS

        self.default_table = QTableWidget(t_rows, t_cols)
        table = self.default_table
        table.setObjectName(CONSTANTS.DEFAULT_TABLE.NAME)
        table.horizontalHeader().setVisible(False)
        table.verticalHeader().setVisible(False)

        for col in range(5):
            table.setColumnWidth(col, CONSTANTS.TABLE_COLUMN_WIDTH.get(col))

        # столбец для ввода
        for row in range(t_rows):
            item = QTableWidgetItem()
            font = QFont('Consolas', 12)
            item.setFont(font)
            table.setItem(row, 3, item)
            if row in CONSTANTS.DEFAULT_TABLE.EDITABLE_ROWS:  # editable
                table.item(row, 3).setBackground(QColor(204, 255, 204))
            else:
                table.item(row, 3).setBackground(QColor(239, 239, 239))
                table.item(row, 3).setFlags(Qt.ItemFlag.ItemIsEnabled)
            table.item(row, 3).setTextAlignment(Qt.AlignmentFlag.AlignCenter)

        for n, row in enumerate((2, 3, 4, 5, 10, 11)):
            table.item(row, 3).setToolTip(f'<span style="font-family: Consolas; color: red">{CONSTANTS.DEFAULT_TABLE.INPUTS_TOOLTIPS[n]}</span>')

        # заголовок
        for row in range(t_rows):
            if row in CONSTANTS.DEFAULT_TABLE.SPAN_ROWS:
                table.setSpan(row, 0, 2, 1)

        for row in range(t_rows):
            if row not in (5, 11):
                self._set_value_in_cell(row, 0, CONSTANTS.DEFAULT_TABLE.HEADER, table)

        for row in range(t_rows):
            self._set_value_in_cell(row, 4, CONSTANTS.DEFAULT_TABLE.UNITS, table)
            table.setRowHeight(row, CONSTANTS.TABLE_ROW_HEIGHT)

        if self.groupbox_count == 1:
            for row in range(t_rows):
                self._set_value_in_cell(row, 1, CONSTANTS.DEFAULT_TABLE.FORMULAS_0, table)
                self._set_value_in_cell(row, 2, CONSTANTS.DEFAULT_TABLE.SYMBOLS_0, table)

        if self.groupbox_count > 1:
            font = QFont('Consolas', 12)
            for row in range(t_rows):
                n = self.groupbox_count
                label = QLabel()
                label.setFont(font)
                label.setAlignment(Qt.AlignmentFlag.AlignCenter)
                label.setStyleSheet('background-color: #EFEFEF')
                table.setCellWidget(row, 2, label)
                if row != t_rows-1:
                    label.setText(CONSTANTS.DEFAULT_TABLE.SYMBOLS_N[row] % (n - 1))
                else:
                    label.setText(CONSTANTS.DEFAULT_TABLE.SYMBOLS_N[row] % n)

                if row != t_rows-1:
                    label = QLabel()
                    label.setFont(font)
                    label.setAlignment(Qt.AlignmentFlag.AlignCenter)
                    label.setStyleSheet('background-color: #EFEFEF')
                    table.setCellWidget(row, 1, label)
                    match row:
                        case 0:
                            if self.groupbox_count == 2:
                                label.setText(CONSTANTS.DEFAULT_TABLE.FORMULAS_N[row] % (n-1, 'n', n-1, n-1, n-1, n-1, n-1))
                            else:
                                label.setText(CONSTANTS.DEFAULT_TABLE.FORMULAS_N[row] % (n-1, str(n-2), n-1, n-1, n-1, n-1, n-1))
                        case 6:
                            label.setText(CONSTANTS.DEFAULT_TABLE.FORMULAS_N[row] % (n-1, n-1, n-1, n-1, n-1))
                        case 7 | 12:
                            label.setText(CONSTANTS.DEFAULT_TABLE.FORMULAS_N[row] % (n-1, n-1, n-1))
                        case 9:
                            label.setText(CONSTANTS.DEFAULT_TABLE.FORMULAS_N[row] % (n-1, n-1, n-1, n-1))
                        case 13:
                            label.setText(CONSTANTS.DEFAULT_TABLE.FORMULAS_N[row] % (n-1, n-1))
                        case _:
                            label.setText(CONSTANTS.DEFAULT_TABLE.FORMULAS_N[row])
                else:
                    label = QLabel()
                    label.setFont(font)
                    label.setAlignment(Qt.AlignmentFlag.AlignCenter)
                    label.setStyleSheet('background-color: #EFEFEF')
                    table.setCellWidget(row, 1, label)
                    label.setText(CONSTANTS.DEFAULT_TABLE.FORMULAS_N[row] % (n, n-1, n-1))

        table.itemChanged.connect(self.validate_input_data_in_default_table)
        table.itemChanged.connect(self.calculate_den_Fn)
        table.itemChanged.connect(self.calculate_Fdpn)
        table.itemChanged.connect(self.calculate_Gn)
        table.itemChanged.connect(self.calculate_Gdpn)
        table.itemChanged.connect(self.calculate_Psn)
        table.itemChanged.connect(self.calculate_Lambda_n)
        table.itemChanged.connect(self.calculate_result)

        table.setStyleSheet('QTableWidget { border: 0px solid grey; font-family: Consolas; }')
        table.setMinimumWidth(1562)
        table.setMinimumHeight(562)
        _layout.addWidget(table)
        _box.setLayout(_layout)
        return _box


    def _set_value_in_cell(self, row, column, data, table) -> None:
        name = table.objectName()
        label = QLabel()
        label.setStyleSheet("background-color: #EFEFEF")
        label.setText(data[row])
        table.setCellWidget(row, column, label)
        if column in (1, 2, 4):
            font = QFont('Consolas', 12)
            label.setFont(font)
            label.setAlignment(Qt.AlignmentFlag.AlignCenter)

            if name == 'table_1' and row in (0, 1, 3, 5, 7) and column == 2:
                label.setToolTip(f'<span style="font-family: Consolas; color: red">{CONSTANTS.TABLE1.VALUES_TOOLTIPS.get(row)}</span>')

            if name == 'table_2' and (row, column) == (3, 2):
                label.setToolTip(f'<span style="font-family: Consolas; color: red">{CONSTANTS.TABLE2.VALUE_TOOLTIP}</span>')


    def _update_num_tables(self) -> None:
        num_tables_label = self.board.itemAtPosition(0, 5).widget()
        num_tables_label.setText(str(self.groupbox_count))


    def _get_default_table_index_nums(self, sender) -> tuple:
        _hash = hash(sender)
        tables = self.findChildren(QTableWidget, CONSTANTS.DEFAULT_TABLE.NAME)
        tables_hashes = [hash(table) for table in tables]
        try:
            index = tables_hashes.index(_hash)
            return (tables[index], index, len(tables))
        except ValueError:
            return (None, None, len(tables))


    def _get_default_table_by_index(self, index) -> object:
        tables = self.findChildren(QTableWidget, CONSTANTS.DEFAULT_TABLE.NAME)
        return tables[index]


    def add_table(self):
        new_table = self.create_default_table()
        self.box_tab1.addWidget(new_table, alignment=Qt.AlignmentFlag.AlignCenter)
        self.box_tab1.addStretch(1)
        self._update_num_tables()


    def delete_table(self) -> None:
        tables = self.findChildren(QTableWidget, CONSTANTS.DEFAULT_TABLE.NAME)
        if len(tables) > 1:
            table_to_remove = tables.pop()
            parent_widget = table_to_remove.parent()
            parent_widget.deleteLater()
            self.groupbox_count -= 1
            self._update_num_tables()
        self._update_result_after_delete_table()


    def copy_table(self) -> None:
        last_table = self.findChildren(QTableWidget, CONSTANTS.DEFAULT_TABLE.NAME)[-1]
        data = [last_table.item(row, 3).text() for row in CONSTANTS.DEFAULT_TABLE.EDITABLE_ROWS]
        if all(data):
            self.add_table()
            last_table = self.findChildren(QTableWidget, CONSTANTS.DEFAULT_TABLE.NAME)[-1]
            for i, row in enumerate(CONSTANTS.DEFAULT_TABLE.EDITABLE_ROWS):
                last_table.item(row, 3).setText(data[i])
        else:
            QMessageBox.information(self, "Информация", 'Чтобы скопировать таблицу, её нужно заполнить.')


    def open(self) -> None:
        options = QFileDialog.Options()
        file_name, _ = QFileDialog.getOpenFileName(self, "Открыть файл", "", "JSON файл (*.json);;Все файлы (*)", options=options)

        tables = self.findChildren(QTableWidget, CONSTANTS.DEFAULT_TABLE.NAME)
        for t in range(1, len(tables)):
            parent_widget = tables[t].parent()
            parent_widget.deleteLater()
            self.groupbox_count -= 1
            self._update_num_tables()

        if file_name:
            try:
                with open(file_name) as f:
                    data = json.load(f)
                    table = self.table_1
                    for n, row in enumerate(CONSTANTS.SAVE_OPEN.TABLE1):
                        table.item(row, 3).setText(data[0]['table_1'][n])

                    self.table_2.item(CONSTANTS.SAVE_OPEN.TABLE2, 3).setText(data[1]['table_2'][0])

                    table = self._get_default_table_by_index(0)
                    for n, row in enumerate(CONSTANTS.SAVE_OPEN.DEFAULT_TABLE):
                        table.item(row, 3).setText(data[2]['default_table'][n])

                    for t in range(3, len(data)):
                        box = self.create_default_table()
                        self.box_tab1.addWidget(box, alignment=Qt.AlignmentFlag.AlignCenter)
                        self.box_tab1.addStretch(1)
                        table = box.findChildren(QTableWidget, CONSTANTS.DEFAULT_TABLE.NAME)[0]
                        for n, row in enumerate(CONSTANTS.SAVE_OPEN.DEFAULT_TABLE):
                            table.item(row, 3).setText(data[t]['default_table'][n])
                    self._update_num_tables()
                self.current_file_path = file_name
                self.setWindowTitle(f'{self.app_title} | {file_name}')

            except Exception as e:
                QMessageBox.critical(self, "Ошибка", f"Не удалось открыть файл: {e}")


    def save(self) -> None:
        if self.current_file_path is None:
            self.save_as()
        else:
            data = self._get_data_for_save()
            with open(self.current_file_path, 'w') as file:
                json.dump(data, file)


    def save_as(self) -> None:
        data = self._get_data_for_save()
        file_name, _ = QFileDialog.getSaveFileName(self, 'Сохранить расчёт', '', 'JSON (*.json)')
        if file_name:
            self.current_file_path = file_name
            self.setWindowTitle(f'{self.app_title} | {file_name}')

            with open(file_name, 'w') as file:
                json.dump(data, file)


    def auto_save(self) -> None:
        if self.current_file_path:
            self.save()
        else:
            self.save_as()


    def export(self) -> None:
        result = self.board.itemAtPosition(0, 3).widget().text()
        if result:
            doc = docx.Document()

            doc_style = doc.styles.add_style('DocStyle', 1)
            doc_style.font.name = 'Times New Roman'
            doc_style.font.size = Pt(12)
            doc.styles['Normal'].base_style = doc_style

            title_style = doc.styles.add_style('TitleStyle', 1)
            title_style.font.name = 'Times New Roman'
            title_style.font.size = Pt(14)
            title_style.font.bold = True
            title = doc.add_paragraph(CONSTANTS.EXPORT.TITLE, style='TitleStyle')
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            header_style = doc.styles.add_style('HeaderStyle', 1)
            header_style.font.bold = True

            # setup fields
            sections = doc.sections
            for section in sections:
                section.left_margin = Cm(2.5)
                section.right_margin = Cm(1)
                section.top_margin = Cm(1.5)
                section.bottom_margin = Cm(1.5)

            # add table 1
            table_1 = doc.add_table(rows=CONSTANTS.TABLE1.ROWS, cols=CONSTANTS.TABLE1.COLUMNS-1)
            doc.add_paragraph()
            table_2 = doc.add_table(rows=CONSTANTS.TABLE2.ROWS, cols=CONSTANTS.TABLE2.COLUMNS-1)
            doc.add_paragraph()

            for t in (table_1, table_2):
                t.alignment = WD_TABLE_ALIGNMENT.CENTER
                t.style = 'Table Grid'

            # setup table header and column widths
            widths = [75, 28, 28, 28]
            header = table_1.rows[0].cells
            for col in range(4):
                header[col].text = CONSTANTS.EXPORT.HEADER[col]
                header[col].width = Mm(widths[col])
                header[col].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                header[col].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            columns = table_2.rows[0].cells
            for col in range(4):
                columns[col].width = Mm(widths[col])
                header[col].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                header[col].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # write headers
            for i in range(1, CONSTANTS.TABLE1.ROWS):
                table_1.columns[0].cells[i].text = CONSTANTS.EXPORT.TABLE1.HEADER[i-1]
                table_1.columns[1].cells[i].text = CONSTANTS.EXPORT.TABLE1.SYMBOLS[i-1]
                table_1.columns[3].cells[i].text = CONSTANTS.EXPORT.TABLE1.UNITS[i-1]
            for i in range(CONSTANTS.TABLE2.ROWS):
                table_2.columns[0].cells[i].text = CONSTANTS.EXPORT.TABLE2.HEADER[i]
                table_2.columns[1].cells[i].text = CONSTANTS.EXPORT.TABLE2.SYMBOLS[i]
                table_2.columns[3].cells[i].text = CONSTANTS.EXPORT.TABLE2.UNITS[i]

            # setup table 1 data
            data = self._get_data_for_export()
            for i in range(1, CONSTANTS.TABLE1.ROWS):
                column = table_1.columns[2]
                column.cells[i].text = data[0]['table_1'][i-1]
            for i in range(CONSTANTS.TABLE2.ROWS):
                column = table_2.columns[2]
                column.cells[i].text = data[1]['table_2'][i]

            # merge table 1 cells
            for j in (2, 4, 6):
                a = table_1.cell(j, 0)
                b = table_1.cell(j+1, 0)
                a.merge(b)
                a.text = a.text.strip()

            # setup alignment
            for row in table_1.rows[1:]:
                for j in range(1, 4):
                    row.cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    row.cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for row in table_2.rows:
                for j in range(1, 4):
                    row.cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    row.cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            default_style = doc.styles.add_style('DefaultStyle', 1)
            default_style.font.name = 'Times New Roman'
            default_style.font.size = Pt(12)
            default_style.font.bold = True

            default_tables = [d for d in data if 'default_table' in d.keys()]
            for k in range(len(default_tables)):
                if k == 0:
                    title = 'Участок до вентилятора'
                else:
                    title = f'Участок {k}'
                default_title = doc.add_paragraph(title, style='DefaultStyle')
                default_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                default_table = doc.add_table(rows=3, cols=4)
                default_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                default_table.style = 'Table Grid'

                columns = default_table.rows[0].cells
                for col in range(4):
                    columns[col].width = Mm(widths[col])
                    header[col].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    header[col].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                for i in range(3):
                    default_table.columns[0].cells[i].text = CONSTANTS.EXPORT.DEFAULT_TABLE.HEADER[i]
                    default_table.columns[3].cells[i].text = CONSTANTS.EXPORT.DEFAULT_TABLE.UNITS[i]

                    n = len(default_tables)
                    if k == 0:
                        default_table.columns[1].cells[i].text = (CONSTANTS.EXPORT.DEFAULT_TABLE.SYMBOLS_0[i])
                    else:
                        if i != 2:
                            default_table.columns[1].cells[i].text = (CONSTANTS.EXPORT.DEFAULT_TABLE.SYMBOLS_N[i] % (n - 1))
                        else:
                            default_table.columns[1].cells[i].text = (CONSTANTS.EXPORT.DEFAULT_TABLE.SYMBOLS_N[i] % n)

                # setup data
                for i in range(3):
                    column = default_table.columns[2]
                    column.cells[i].text = default_tables[k]['default_table'][i]

                # setup alignment
                for row in default_table.rows:
                    for j in range(1, 4):
                        row.cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        row.cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            result_style = doc.styles.add_style('ResultStyle', 1)
            result_style.font.name = 'Times New Roman'
            result_style.font.size = Pt(12)
            result_style.font.bold = True
            result = doc.add_paragraph(CONSTANTS.EXPORT.RESULT.TITLE, style='ResultStyle')
            result.alignment = WD_ALIGN_PARAGRAPH.CENTER

            result_table = doc.add_table(rows=1, cols=4)
            result_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            result_table.style = 'Table Grid'

            widths = [68, 38.5, 38.5, 38.5]
            header = result_table.rows[0].cells
            for col in range(4):
                if col == 2:
                    header[col].text = data[-1]['result']
                else:
                    header[col].text = CONSTANTS.EXPORT.RESULT.DATA[col]
                if col != 0:
                    header[col].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    header[col].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                header[col].width = Mm(widths[col])

            now = datetime.now()
            date_time = now.strftime("%d_%m_%y_%H_%M")
            file_name = f'{self.current_file_path}_{date_time}'
            file_name = file_name.replace('.json', '')
            doc.save(f'{file_name}.docx')
            QMessageBox.information(self, 'Информация', 'Расчёт успешно экспортирован')
            os.startfile(os.path.join(basedir, f'{file_name}.docx'))
        else:
            QMessageBox.critical(self, 'Ошибка', 'Пока что нечего экспортировать')


    def _get_data_for_save(self) -> list:
        data = []
        tables = self.findChildren(QTableWidget)

        for table in tables:
            match table.objectName():
                case 'table_1':
                    table_data = [table.item(row, 3).text() for row in CONSTANTS.SAVE_OPEN.TABLE1]
                    data.append({'table_1': table_data})
                case 'table_2':
                    table_data = [table.item(CONSTANTS.SAVE_OPEN.TABLE2, 3).text()]
                    data.append({'table_2': table_data})
                case 'default_table':
                    table_data = [table.item(row, 3).text() for row in CONSTANTS.SAVE_OPEN.DEFAULT_TABLE]
                    data.append({'default_table': table_data})
        return data


    def _get_data_for_export(self) -> list:
        data = []
        tables = self.findChildren(QTableWidget)
        for table in tables:
            match table.objectName():
                case 'table_1':
                    table_data = [table.item(row, 3).text() for row in range(CONSTANTS.TABLE1.ROWS)]
                    data.append({'table_1': table_data})
                case 'table_2':
                    table_data = [table.item(row, 3).text() for row in range(CONSTANTS.TABLE2.ROWS)]
                    data.append({'table_2': table_data})
                case 'default_table':
                    table_data = [table.item(row, 3).text() for row in CONSTANTS.DEFAULT_TABLE.EXPORT_ROWS]
                    data.append({'default_table': table_data})
        result = self.board.itemAtPosition(0, 3).widget().text()
        data.append({'result': result})
        return data


    def closeEvent(self, event) -> None:
        if self.current_file_path is None:
            reply = QMessageBox.question(
                self,
                'Подтверждение',
                '''<html>Вы уверены, что хотите закрыть программу?<br><font color="red">Не сохраненный расчет будет потерян.</font></html>
                ''',
                QMessageBox.Yes | QMessageBox.No, QMessageBox.No
            )

            if reply == QMessageBox.No:
                self.save()
                event.accept()
                event.ignore()
            else:
                event.accept()
        else:
            reply = QMessageBox.question(self, 'Подтверждение', 'Сохранить текущие изменения?', QMessageBox.Yes | QMessageBox.No, QMessageBox.No)

            if reply == QMessageBox.No:
                event.accept()
            else:
                self.save
                event.accept()


    def open_manual(self) -> None:
        if platform.system() == "Windows":
            os.startfile(os.path.join(basedir, 'recalculation_smoke_exhaust_fan_manual.pdf'))


    def show_about(self) -> None:
        QMessageBox.information(self, "О программе", CONSTANTS.ABOUT)


    def validate_input_data_in_tables(self, item) -> None:
        table = self.sender().objectName()
        row = item.row()
        if (table == 'table_1' and row in (0, 1, 3, 5, 7)) or table == 'table_2' and row == 3:
            value = (table, row)

            # table_1
            table_1_pattern_0 = r'^(?:[0-9]|[1-9][0-9]{1,2}|1[0-9]{3}|2000|2500)$'  # 0...2_500
            table_1_pattern_1_5 = r'^[0-9]{1,3}$|^1000$'  # 0...1_000
            table_1_pattern_3 = r'^-?[0-4]?[0-9]$|^50$'  # -50...50
            table_1_pattern_7 = r'^(?:\d{1,2}|1\d{2}|2[0-9]{2}|300)(?:\.\d{1,2})?$'  # 0...300.00

            # table_2
            table_2_pattern_3 = r'^(?:[0-9]|[1-9][0-9]{1,4}|100000)$'  # 0...100_000

            match value:
                case 'table_1', 0:
                    pattern = table_1_pattern_0
                case 'table_1', 1:
                    pattern = table_1_pattern_1_5
                case 'table_1', 5:
                    pattern = table_1_pattern_1_5
                case 'table_1', 3:
                    pattern = table_1_pattern_3
                case 'table_1', 7:
                    pattern = table_1_pattern_7
                case 'table_2', 3:
                    pattern = table_2_pattern_3

            if not re.match(pattern, item.text()):
                item.setText('')
            else:
                item.setText(item.text())


    def validate_input_data_in_default_table(self, item) -> None:
        row = item.row()
        if row in (1, 2, 3, 4, 5, 10, 11):
            pattern_1_2 = r'^(?:[0-9]|[1-9]\d|100)(?:\.\d{1,3})?$'  # 0...100.00
            pattern_3 = r'^(?:\d|[1][0-9]|20)(?:\.\d{1,2})?$'  # 0...20.00
            pattern_4_10 = r'^(?!0(?:\.0+)?$)(?:[0-4](?:\.\d)?|5(?:\.0)?)$'  # 0.1...5.0
            pattern_5_11 = r'^(?:0\.[1-9]|[1-5](?:\.\d)?|6(?:\.0)?)$'  # 0.1...6.0

            match row:
                case 1 | 2:
                    pattern = pattern_1_2
                case 3:
                    pattern = pattern_3
                case 4 | 10:
                    pattern = pattern_4_10
                case 5 | 11:
                    pattern = pattern_5_11

            if not re.match(pattern, item.text()):
                item.setText('')
            else:
                item.setText(item.text())


    def calculate_Tsm0(self, row, column) -> None:
        if row == 1:
            table = self.table_1
            temperature = table.item(1, 3).text()
            if temperature:
                temperature = int(temperature)
                result = "{:.2f}".format(round(273.15 + temperature, 2))
                table.item(2, 3).setText(result)
            else:
                table.item(2, 3).setText('')


    def calculate_Ta(self, row, column) -> None:
        if row == 3:
            table = self.table_1
            temperature = table.item(3, 3).text()
            if temperature:
                temperature = int(temperature)
                result = "{:.2f}".format(round(273.15 + temperature, 2))
                table.item(4, 3).setText(result)
            else:
                table.item(4, 3).setText('')


    def calculate_Tv(self, row, column) -> None:
        if row == 5:
            table = self.table_1
            temperature = table.item(5, 3).text()
            if temperature:
                temperature = int(temperature)
                result = "{:.2f}".format(round(273.15 + temperature, 2))
                table.item(6, 3).setText(result)
            else:
                table.item(6, 3).setText('')


    def calculate_density_a(self, row, column) -> None:
        if row == 4:
            table = self.table_1
            temperature = table.item(4, 3).text()
            if temperature:
                temperature = float(temperature)
                result = "{:.3f}".format(round(353 / temperature, 3))
                table.item(8, 3).setText(result)
            else:
                table.item(8, 3).setText('')


    def calculate_density_v(self, row, column) -> None:
        if row == 6:
            table = self.table_1
            temperature = table.item(6, 3).text()
            if temperature:
                temperature = float(temperature)
                result = "{:.3f}".format(round(353 / temperature, 3))
                table.item(9, 3).setText(result)
            else:
                table.item(9, 3).setText('')


    def calculate_density_sm(self, row, column) -> None:
        if row in (2, 4, 6, 8):
            table_1 = self.table_1
            table_2 = self.table_2
            temperature_sm0 = table_1.item(2, 3).text()
            temperature_a = table_1.item(4, 3).text()
            temperature_v = table_1.item(6, 3).text()
            density_a = table_1.item(8, 3).text()
            if all([temperature_sm0, temperature_a, temperature_v, density_a]):
                temperature_sm0 = float(temperature_sm0)
                temperature_a = float(temperature_a)
                temperature_v = float(temperature_v)
                density_a = float(density_a)
                result = 2 * density_a * temperature_a / (temperature_sm0 + temperature_v)
                result = "{:.3f}".format(round(result, 3))
                table_2.item(0, 3).setText(result)
            else:
                table_2.item(0, 3).setText('')


    def calculate_Psa(self, row, column) -> None:
        table = self.sender().objectName()
        if (table == 'table_1' and row in (0, 7, 8, 9)) or (table == 'table_2' and row == 0):
            table_1 = self.table_1
            table_2 = self.table_2
            pressure = table_1.item(0, 3).text()
            height = table_1.item(7, 3).text()
            density_a = table_1.item(8, 3).text()
            density_v = table_1.item(9, 3).text()
            density_sm = table_2.item(0, 3).text()
            if all([pressure, height, density_a, density_v, density_sm]):
                pressure = float(pressure)
                height = float(height)
                density_a = float(density_a)
                density_v = float(density_v)
                density_sm = float(density_sm)
                result = pressure * density_v / 1.2 + 9.81 * height * (density_a - density_sm)
                result = "{:.0f}".format(round(result, 0))
                table_2.item(1, 3).setText(result)
            else:
                table_2.item(1, 3).setText('')


    def calculate_pressure_after_Psa(self, row, column) -> None:
        table = self.sender().objectName()
        if (table == 'table_1' and row == 9) or (table == 'table_2' and row == 1):
            density_v = self.table_1.item(9, 3).text()
            pressure_sa = self.table_2.item(1, 3).text()
            if all([density_v, pressure_sa]):
                density_v = float(density_v)
                pressure_sa = float(pressure_sa)
                result = pressure_sa * 1.2 / density_v
                result = "{:.0f}".format(round(result, 0))
                self.table_2.item(2, 3).setText(result)
            else:
                self.table_2.item(2, 3).setText('')


    def calculate_Ga(self, row, column) -> None:
        table = self.sender().objectName()
        if (table == 'table_1' and row == 8) or (table == 'table_2' and row == 3):
            density_a = self.table_1.item(8, 3).text()
            flow_a = self.table_2.item(3, 3).text()
            if all([density_a, flow_a]):
                density_a = float(density_a)
                flow_a = float(flow_a)
                result = flow_a * density_a / 3_600
                result = "{:.3f}".format(round(result, 3))
                self.table_2.item(4, 3).setText(result)
            else:
                self.table_2.item(4, 3).setText('')


    def calculate_den_Fn(self, item) -> None:
        if item.row() in (4, 5):
            sender = self.sender()
            table = self._get_default_table_index_nums(sender)[0]

            a_n = table.item(4, 3).text()
            b_n = table.item(5, 3).text()
            if all([a_n, b_n]):
                a_n = float(a_n)
                b_n = float(b_n)
                result_1 = 2 * a_n * b_n / (a_n + b_n)
                result_1 = "{:.3f}".format(round(result_1, 3))
                table.item(6, 3).setText(result_1)
                result_2 = "{:.2f}".format(round(a_n * b_n, 2))
                table.item(7, 3).setText(result_2)
            else:
                table.item(6, 3).setText('')
                table.item(7, 3).setText('')


    def calculate_Fdpn(self, item) -> None:
        if item.row() in (10, 11):
            sender = self.sender()
            table = self._get_default_table_index_nums(sender)[0]

            a_dpn = table.item(10, 3).text()
            b_dpn = table.item(11, 3).text()
            if all([a_dpn, b_dpn]):
                a_dpn = float(a_dpn)
                b_dpn = float(b_dpn)
                result = a_dpn * b_dpn
                result = "{:.2f}".format(round(result, 2))
                table.item(12, 3).setText(result)
            else:
                table.item(12, 3).setText('')


    def calculate_Gn(self, item) -> None:
        sender = self.sender()
        sender_name = sender.objectName()
        row = item.row()
        if sender_name == 'table_2' and row == 4:
            G_a = self.table_2.item(4, 3).text()
            table = self._get_default_table_by_index(0)
            G_dpn = table.item(9, 3).text()
            if all([G_a, G_dpn]):
                G_a = float(G_a)
                G_dpn = float(G_dpn)
                result = G_a - G_dpn
                result = "{:.3f}".format(round(result, 3))
                table.item(13, 3).setText(result)
            else:
                table.item(13, 3).setText('')

        elif self._get_default_table_index_nums(sender)[1] == 0 and row == 9:
            G_a = self.table_2.item(4, 3).text()
            table = self._get_default_table_by_index(0)
            G_dpn = table.item(9, 3).text()
            if all([G_a, G_dpn]):
                G_a = float(G_a)
                G_dpn = float(G_dpn)
                result = G_a - G_dpn
                result = "{:.3f}".format(round(result, 3))
                table.item(13, 3).setText(result)
            else:
                table.item(13, 3).setText('')

        elif self._get_default_table_index_nums(sender)[1] != 0 and row == 9:
            tables = self.findChildren(QTableWidget, CONSTANTS.DEFAULT_TABLE.NAME)
            for i in range(1, len(tables)):
                G_1 = tables[i-1].item(13, 3).text()
                G_dpn = tables[i].item(9, 3).text()
                if all([G_1, G_dpn]):
                    G_1 = float(G_1)
                    G_dpn = float(G_dpn)
                    result = G_1 - G_dpn
                    result = "{:.3f}".format(round(result, 3))
                    tables[i].item(13, 3).setText(result)
                else:
                    tables[i].item(13, 3).setText('')

        elif sender_name == 'default_table' and row == 13:
            tables = self.findChildren(QTableWidget, CONSTANTS.DEFAULT_TABLE.NAME)
            table, index, nums = self._get_default_table_index_nums(sender)
            for i in range(index+1, nums):
                G_1 = tables[i-1].item(13, 3).text()
                G_dpn = tables[i].item(9, 3).text()
                if all([G_1, G_dpn]):
                    G_1 = float(G_1)
                    G_dpn = float(G_dpn)
                    result = G_1 - G_dpn
                    result = "{:.3f}".format(round(result, 3))
                    tables[i].item(13, 3).setText(result)
                else:
                    tables[i].item(13, 3).setText('')


    def calculate_Gdpn(self, item) -> None:
        sender = self.sender()
        sender_name = sender.objectName()
        row = item.row()

        if (sender_name == 'table_1' and row == 3):
            tables = self.findChildren(QTableWidget, CONSTANTS.DEFAULT_TABLE.NAME)
            for i in range(len(tables)):
                temperature_a = self.table_1.item(3, 3).text()
                Psn = tables[i].item(0, 3).text()
                Fdpn = tables[i].item(12, 3).text()
                if all([temperature_a, Psn, Fdpn]):
                    temperature_a = float(temperature_a)
                    devider = 29_900 if temperature_a <= 20 else 30_300
                    Psn = float(Psn)
                    Fdpn = float(Fdpn)
                    result = Fdpn * pow(Psn / devider, 0.5)
                    result = "{:.3f}".format(round(result, 3))
                    tables[i].item(9, 3).setText(result)
                else:
                    tables[i].item(9, 3).setText('')

        elif sender_name == 'default_table' and row in (0, 12):
            table = self._get_default_table_index_nums(sender)[0]
            temperature_a = self.table_1.item(3, 3).text()
            Psn = table.item(0, 3).text()
            Fdpn = table.item(12, 3).text()
            if all([temperature_a, Psn, Fdpn]):
                temperature_a = float(temperature_a)
                devider = 29_900 if temperature_a <= 20 else 30_300
                Psn = float(Psn)
                Fdpn = float(Fdpn)
                result = Fdpn * pow(Psn / devider, 0.5)
                result = "{:.3f}".format(round(result, 3))
                table.item(9, 3).setText(result)
            else:
                table.item(9, 3).setText('')


    def calculate_Psn(self, item) -> None:
        sender = self.sender()
        sender_name = sender.objectName()
        row = item.row()

        if sender_name == 'table_2' and row in (1, 4):
            table = self._get_default_table_by_index(0)
            density_a = self.table_1.item(8, 3).text()
            Psa = self.table_2.item(1, 3).text()
            Ga = self.table_2.item(4, 3).text()
            Ln = table.item(1, 3).text()
            KMSn = table.item(2, 3).text()
            den = table.item(6, 3).text()
            Fn = table.item(7, 3).text()
            Lambda_n = table.item(8, 3).text()
            if all([density_a, Psa, Ga, Ln, KMSn, den, Fn, Lambda_n]):
                result = self._calculate_Psn_get_result(density_a, Psa, Ga, Ln, KMSn, den, Fn, Lambda_n)
                table.item(0, 3).setText(result)
            else:
                table.item(0, 3).setText('')

        elif sender_name == 'table_1' and row == 8:
            tables = self.findChildren(QTableWidget, CONSTANTS.DEFAULT_TABLE.NAME)
            for i in range(1, len(tables)):
                density_a = self.table_1.item(8, 3).text()
                Psn = tables[i-1].item(1, 3).text()
                Gn = tables[i-1].item(4, 3).text()
                Ln = tables[i].item(1, 3).text()
                KMSn = tables[i].item(2, 3).text()
                den = tables[i].item(6, 3).text()
                Fn = tables[i].item(7, 3).text()
                Lambda_n = tables[i].item(8, 3).text()
                if all([density_a, Psn, Gn, Ln, KMSn, den, Fn, Lambda_n]):
                    result = self._calculate_Psn_get_result(density_a, Psn, Gn, Ln, KMSn, den, Fn, Lambda_n)
                    tables[i].item(0, 3).setText(result)
                else:
                    tables[i].item(0, 3).setText('')

        elif sender_name == 'default_table' and row in (1, 2, 6, 7, 8):
            table, index, nums = self._get_default_table_index_nums(sender)
            density_a = self.table_1.item(8, 3).text()

            if self._get_default_table_index_nums(sender)[1] == 0:
                Psn = self.table_2.item(1, 3).text()
                Gn = self.table_2.item(4, 3).text()
            else:
                Psn = self._get_default_table_by_index(index-1).item(0, 3).text()
                Gn = self._get_default_table_by_index(index-1).item(13, 3).text()

            Ln = table.item(1, 3).text()
            KMSn = table.item(2, 3).text()
            den = table.item(6, 3).text()
            Fn = table.item(7, 3).text()
            Lambda_n = table.item(8, 3).text()
            if all([density_a, Psn, Gn, Ln, KMSn, den, Fn, Lambda_n]):
                result = self._calculate_Psn_get_result(density_a, Psn, Gn, Ln, KMSn, den, Fn, Lambda_n)
                table.item(0, 3).setText(result)
            else:
                table.item(0, 3).setText('')

        elif sender_name == 'default_table' and row == 1:
            tables = self.findChildren(QTableWidget, CONSTANTS.DEFAULT_TABLE.NAME)
            table, index, nums = self._get_default_table_index_nums(sender)
            for i in range(index+1, nums):
                density_a = self.table_1.item(8, 3).text()
                Psn = tables[i-1].item(1, 3).text()
                Gn = tables[i-1].item(4, 3).text()
                Ln = tables[i].item(1, 3).text()
                KMSn = tables[i].item(2, 3).text()
                den = tables[i].item(6, 3).text()
                Fn = tables[i].item(7, 3).text()
                Lambda_n = tables[i].item(8, 3).text()
                if all([density_a, Psn, Gn, Ln, KMSn, den, Fn, Lambda_n]):
                    result = self._calculate_Psn_get_result(density_a, Psn, Gn, Ln, KMSn, den, Fn, Lambda_n)
                    tables[i].item(0, 3).setText(result)
                else:
                    tables[i].item(0, 3).setText('')


    def _calculate_Psn_get_result(self, density, Ps, G, L, KMS, de, F, Lambda) -> str:
        density_a = float(density)
        Psa = float(Ps)
        Ga = float(G)
        Ln = float(L)
        KMSn = float(KMS)
        den = float(de)
        Fn = float(F)
        Lambda_n = float(Lambda)
        result = Psa - 0.5 * density_a * (KMSn + Lambda_n * Ln / den) * pow((Ga / (density_a * Fn)), 2)
        result = "{:.0f}".format(round(result, 0))
        return result


    def calculate_Lambda_n(self, item) -> None:
        sender = self.sender()
        sender_name = sender.objectName()
        row = item.row()

        if sender_name == 'table_2' and row == 4:
            Ta = self.table_1.item(4, 3).text()
            density_a = self.table_1.item(8, 3).text()
            G_a = self.table_2.item(4, 3).text()
            table = self._get_default_table_by_index(0)
            kn = table.item(3, 3).text()
            den = table.item(6, 3).text()
            Fn = table.item(7, 3).text()
            if all([Ta, density_a, G_a, kn, den, Fn]):
                Ta = float(Ta)
                density_a = float(density_a)
                G_a = float(G_a)
                kn = float(kn)
                den = float(den)
                Fn = float(Fn)
                result = 0.11 * pow(68 / (density_a * (G_a / density_a / Fn) * den / ((pow(-10, -5) * pow((Ta - 273.15), 2) + 0.0419 * (Ta - 273.15) + 16.229) / 1_000_000)) + kn / den, 0.25)
                result = "{:.4f}".format(round(result, 4))
                table.item(8, 3).setText(result)
            else:
                table.item(8, 3).setText('')

        elif sender_name == 'table_1' and row in (4, 8):
            tables = self.findChildren(QTableWidget, CONSTANTS.DEFAULT_TABLE.NAME)
            for i in range(len(tables)):
                Ta = self.table_1.item(4, 3).text()
                density_a = self.table_1.item(8, 3).text()

                if self._get_default_table_index_nums(sender)[1] == 0:
                    G_n = self.table_2.item(4, 3).text()
                else:
                    G_n = tables[i-1].item(13, 3).text()

                kn = tables[i].item(3, 3).text()
                den = tables[i].item(6, 3).text()
                Fn = tables[i].item(7, 3).text()
                if all([Ta, density_a, G_n, kn, den, Fn]):
                    Ta = float(Ta)
                    density_a = float(density_a)
                    G_n = float(G_n)
                    kn = float(kn)
                    den = float(den)
                    Fn = float(Fn)
                    result = 0.11 * pow(68 / (density_a * (G_n / density_a / Fn) * den / ((pow(-10, -5) * pow((Ta - 273.15), 2) + 0.0419 * (Ta - 273.15) + 16.229) / 1_000_000)) + kn / den, 0.25)
                    result = "{:.4f}".format(round(result, 4))
                    tables[i].item(8, 3).setText(result)
                else:
                    tables[i].item(8, 3).setText('')

        elif sender_name == 'default_table' and row in (3, 6, 7):
            table, index, nums = self._get_default_table_index_nums(sender)
            Ta = self.table_1.item(4, 3).text()
            density_a = self.table_1.item(8, 3).text()
            kn = table.item(3, 3).text()
            den = table.item(6, 3).text()
            Fn = table.item(7, 3).text()

            if self._get_default_table_index_nums(sender)[1] == 0:
                G_n = self.table_2.item(4, 3).text()
            else:
                G_n = self._get_default_table_by_index(index-1).item(13, 3).text()

            if all([Ta, density_a, G_n, kn, den, Fn]):
                Ta = float(Ta)
                density_a = float(density_a)
                G_n = float(G_n)
                kn = float(kn)
                den = float(den)
                Fn = float(Fn)
                result = 0.11 * pow(68 / (density_a * (G_n / density_a / Fn) * den / ((pow(-10, -5) * pow((Ta - 273.15), 2) + 0.0419 * (Ta - 273.15) + 16.229) / 1_000_000)) + kn / den, 0.25)
                result = "{:.4f}".format(round(result, 4))
                table.item(8, 3).setText(result)
            else:
                table.item(8, 3).setText('')

        elif sender_name == 'default_table' and row == 13:
            tables = self.findChildren(QTableWidget, CONSTANTS.DEFAULT_TABLE.NAME)
            table, index, nums = self._get_default_table_index_nums(sender)
            for i in range(index+1, nums):
                Ta = self.table_1.item(4, 3).text()
                density_a = self.table_1.item(8, 3).text()
                table = self._get_default_table_by_index(0)
                G_n = tables[i-1].item(13, 3).text()
                kn = tables[i].item(3, 3).text()
                den = tables[i].item(6, 3).text()
                Fn = tables[i].item(7, 3).text()
                if all([Ta, density_a, G_n, kn, den, Fn]):
                    Ta = float(Ta)
                    density_a = float(density_a)
                    G_n = float(G_n)
                    kn = float(kn)
                    den = float(den)
                    Fn = float(Fn)
                    result = 0.11 * pow(68 / (density_a * (G_n / density_a / Fn) * den / ((pow(-10, -5) * pow((Ta - 273.15), 2) + 0.0419 * (Ta - 273.15) + 16.229) / 1_000_000)) + kn / den, 0.25)
                    result = "{:.4f}".format(round(result, 4))
                    tables[i].item(8, 3).setText(result)
                else:
                    tables[i].item(8, 3).setText('')


    def calculate_result(self, item) -> None:
        sender = self.sender()
        sender_name = sender.objectName()
        row = item.row()

        if (sender_name == 'table_1' and row == 8) or (sender_name == 'table_2' and row == 4) or (sender_name == 'default_table' and row == 13):
            self._set_result()


    def _update_result_after_delete_table(self) -> None:
        self._set_result()


    def _set_result(self) -> None:
        density_a = self.table_1.item(8, 3).text()
        last_table = self.findChildren(QTableWidget, CONSTANTS.DEFAULT_TABLE.NAME)[self.groupbox_count-1]
        G_last = last_table.item(13, 3).text()
        if all([density_a, G_last]):
            density_a = float(density_a)
            G_last = float(G_last)
            result = 3_600 * G_last / density_a
            result = "{:.0f}".format(round(result, 0))
            self.board.itemAtPosition(0, 3).widget().setText(result)
        else:
            self.board.itemAtPosition(0, 3).widget().setText('')


    def _set_Ta_in_board(self, row, column) -> None:
        if (row, column) == (4, 3):
            Ta = self.table_1.item(4, 3).text()
            if Ta:
                self.board.itemAtPosition(0, 1).widget().setText(Ta)
            else:
                self.board.itemAtPosition(0, 1).widget().setText('-')







if __name__ == '__main__':
    app = QApplication(sys.argv)
    app.setFont(QFont('Consolas', 10))
    app.setStyleSheet('QMessageBox { messagebox-text-interaction-flags: 5; font-size: 13px; }')
    app.setStyle('windowsvista')
    window = MainWindow()
    window.setWindowIcon(QIcon(os.path.join(basedir, 'app.ico')))
    window.setIconSize(QSize(15, 15))
    window.show()
    sys.exit(app.exec())
